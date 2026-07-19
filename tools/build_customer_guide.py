from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "HƯỚNG DẪN SỬ DỤNG HỆ THỐNG NHÓM HỌC TẬP - HVHN.docx"
LOGO = ROOT / "hvn.jpg"

INK = "17324D"
BLUE = "1F4D78"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F6F9FC"
MUTED = "526777"
LINE = "C9D6E2"


def set_run_font(run, *, size, color=INK, bold=False, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(width_dxa / 1440)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=LINE):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(paragraph, text, *, size=11, color=INK, bold=False, italic=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return run


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relation_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    fonts.set(qn("w:eastAsia"), "Arial")
    r_pr.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    r_pr.append(size)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    add_text(paragraph, text, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else INK, bold=True)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    add_text(paragraph, text)
    return paragraph


def add_numbered_step(doc, title, detail):
    paragraph = doc.add_paragraph(style="List Number")
    add_text(paragraph, title + ". ", bold=True)
    add_text(paragraph, detail)
    return paragraph


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="B8CBDD")
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=3, line=1.2)
    add_text(paragraph, title, size=11, color=BLUE, bold=True)
    paragraph = cell.add_paragraph()
    set_paragraph_spacing(paragraph, after=0, line=1.2)
    add_text(paragraph, body, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, INK),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_and_footer(section):
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    add_text(paragraph, "HVHN  |  HƯỚNG DẪN THÀNH VIÊN", size=8.5, color=MUTED, bold=True)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, before=0, after=0, line=1.0)
    add_text(paragraph, "Hồn Văn · Hồn Người  |  Trang ", size=8.5, color=MUTED)
    add_page_number(paragraph)


def add_channel_map(doc):
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    set_table_borders(table)
    header = table.rows[0].cells
    for cell, label in zip(header, ("Kênh", "Dùng để")):
        set_cell_shading(cell, BLUE)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        set_paragraph_spacing(paragraph, after=0, line=1.1)
        add_text(paragraph, label, size=10.5, color="FFFFFF", bold=True)
    rows = (
        ("#hỏi-đáp-bài-tập", "Đặt câu hỏi học tập; mỗi vấn đề nên có một bài hoặc thread riêng."),
        ("#chia-sẻ-tài-liệu", "Chia sẻ đề, dàn ý và tài liệu học tập phù hợp."),
        ("#thảo-luận-văn-học", "Trao đổi cảm nhận, tác phẩm và vấn đề văn học."),
        ("#thư-viện-tài-liệu", "Tra cứu tài liệu dùng chung."),
        ("#lệnh-bot-chung", "Sử dụng các lệnh của Then, tránh làm loãng kênh học tập."),
        ("#dịch-vụ-ngoài", "Xem thông tin dịch vụ hỗ trợ chuyên sâu ngoài chương trình."),
    )
    for row_index, values in enumerate(rows, 1):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_shading(cell, PALE_BLUE if row_index % 2 else "FFFFFF")
            cell.text = ""
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line=1.15)
            add_text(paragraph, value, size=10.3, color=INK, bold=(cell is cells[0]))
    doc.add_paragraph()


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    set_styles(doc)
    add_header_and_footer(section)
    doc.core_properties.title = "Hướng dẫn sử dụng hệ thống nhóm học tập HVHN"
    doc.core_properties.subject = "Hướng dẫn dành cho thành viên mới"
    doc.core_properties.author = "Hồn Văn · Hồn Người"

    if LOGO.exists():
        logo = doc.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(logo, after=5, line=1.0)
        logo.add_run().add_picture(str(LOGO), width=Inches(0.78))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, before=0, after=4, line=1.0)
    add_text(title, "HƯỚNG DẪN THAM GIA VÀ SỬ DỤNG", size=21, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=14, line=1.0)
    add_text(subtitle, "HỆ THỐNG NHÓM HỌC TẬP HỒN VĂN · HỒN NGƯỜI", size=11, color=BLUE, bold=True)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(info, after=14, line=1.15)
    add_text(info, "Nhóm Zalo: ", size=10.5, color=MUTED, bold=True)
    add_hyperlink(info, "https://zalo.me/g/wpig3zpoz9rxdj1cjena", "https://zalo.me/g/wpig3zpoz9rxdj1cjena")

    add_callout(
        doc,
        "Bắt đầu đúng cách",
        "Hãy hoàn thành bốn bước dưới đây theo thứ tự. Chỉ mất vài phút, nhưng giúp quyền truy cập, học liệu và Then hoạt động đúng với tài khoản của bạn.",
    )

    add_heading(doc, "1. Bắt đầu trong 4 bước")
    add_numbered_step(doc, "Đọc thông tin chung", "Vào #sảnh-chào-mừng để xem thông báo đầu tiên, sau đó đọc đầy đủ #luật-lệ.")
    add_numbered_step(doc, "Xác nhận thành viên", "Tại #cổng-xác-nhận, nhấn nút xác nhận để mở khóa các kênh học tập.")
    add_numbered_step(doc, "Đọc hướng dẫn Then", "Đọc kỹ #hướng-dẫn-dùng-bot và bấm xác nhận trước khi dùng AI. Đây là bước bắt buộc để Then hiểu đúng yêu cầu của bạn hơn.")
    add_numbered_step(doc, "Kích hoạt học liệu và Then trên web", "Nếu bạn nhận học liệu qua HVHN, vào #truy-cập-tài-liệu, điền đúng Họ và tên cùng Email. Sau đó đăng nhập Google bằng chính email đã khai báo.")

    add_heading(doc, "2. Sử dụng Then hiệu quả")
    add_heading(doc, "Trên Discord", level=2)
    paragraph = doc.add_paragraph()
    add_text(paragraph, "Dùng Then tại #lệnh-bot-chung. Các lệnh thường dùng gồm: ", bold=False)
    add_text(paragraph, "/ai", color=BLUE, bold=True)
    add_text(paragraph, " (hỏi trợ giảng AI), ", color=INK)
    add_text(paragraph, "/van_hoi", color=BLUE, bold=True)
    add_text(paragraph, " (hỏi về bài Văn, luận điểm, dẫn chứng), ", color=INK)
    add_text(paragraph, "/goi_y_mo_bai", color=BLUE, bold=True)
    add_text(paragraph, " và ", color=INK)
    add_text(paragraph, "/luyen_de_hom_nay", color=BLUE, bold=True)
    add_text(paragraph, ".", color=INK)
    add_bullet(doc, "Nêu rõ đề bài, phần bạn cần hỗ trợ và kết quả mong muốn. Khi hỏi về một đoạn trích, hãy gửi kèm đoạn trích hoặc dữ liệu cần phân tích.")
    add_bullet(doc, "Nếu hỏi nối tiếp, hãy dán lại thông tin quan trọng của câu trước để Then có đủ ngữ cảnh.")
    add_bullet(doc, "Mỗi thành viên có tối đa 7 lượt AI trong 24 giờ và 30 lượt trong 7 ngày; chỉ lượt trả lời thành công mới được tính.")

    add_heading(doc, "Then trên website", level=2)
    add_numbered_step(doc, "Mở đường dẫn Then", "Tìm liên kết tại #bảng-tin-thông-báo, #hướng-dẫn-dùng-bot hoặc #truy-cập-tài-liệu.")
    add_numbered_step(doc, "Khởi động phiên làm việc", "Tại website, chọn Start (Bắt đầu), rồi nhập yêu cầu vào khung chat.")
    add_numbered_step(doc, "Bổ sung ngữ liệu khi cần", "Khi giao diện hỗ trợ, bạn có thể đính kèm PDF, MD, ảnh sơ đồ tư duy hoặc liên kết YouTube để Then có thêm căn cứ xử lý yêu cầu.")

    add_callout(
        doc,
        "Lưu ý về AI",
        "Then là công cụ hỗ trợ tư duy, không thay thế việc học và kiểm chứng. Luôn đối chiếu kiến thức với tác phẩm gốc, SGK hoặc nguồn đáng tin cậy trước khi sử dụng trong bài làm.",
    )

    add_heading(doc, "3. Bản đồ các kênh cần biết")
    add_channel_map(doc)

    add_heading(doc, "4. Học liệu và quyền truy cập")
    add_bullet(doc, "Chỉ dùng một tài khoản và một email trong suốt quá trình tham gia để hệ thống đồng bộ chính xác.")
    add_bullet(doc, "Hệ thống hỗ trợ đăng nhập tối đa 3 thiết bị. Nếu cần hỗ trợ khi quyền truy cập bị tạm khóa, hãy liên hệ quản trị viên.")
    add_bullet(doc, "Tài liệu được gắn watermark theo Họ và tên cùng Email người nhận; không hỗ trợ tải xuống hoặc in ấn. Quyền truy cập có thời hạn 30 ngày kể từ ngày được cấp.")

    add_heading(doc, "5. Dịch vụ hỗ trợ chuyên sâu ngoài chương trình")
    paragraph = doc.add_paragraph()
    add_text(paragraph, "Khi cần hỗ trợ chuyên sâu hơn chatbot, bạn có thể xem thông báo tại #dịch-vụ-ngoài hoặc liên hệ quản trị viên. Các dịch vụ gồm:")
    add_bullet(doc, "Chấm chữa bài chi tiết: góp ý lập luận, dẫn chứng, diễn đạt và hướng nâng cấp bài viết.")
    add_bullet(doc, "Biên soạn bài mẫu theo yêu cầu: bài tham khảo, mở bài hoặc kết bài theo đề và trọng tâm bạn cần luyện.")
    add_bullet(doc, "Thiết kế đề thi độc quyền: đề luyện theo cấu trúc phù hợp với mục tiêu học tập hoặc bồi dưỡng học sinh giỏi.")
    add_callout(
        doc,
        "Khi đăng ký dịch vụ",
        "Hãy nêu rõ mục tiêu sử dụng, yêu cầu cụ thể và thời hạn cần nhận. Đội ngũ sẽ trao đổi phạm vi, thời gian và chi phí trước khi nhận yêu cầu.",
    )

    add_heading(doc, "6. Khi cần hỗ trợ")
    add_bullet(doc, "Nếu câu trả lời của Then trên Discord chưa phù hợp, dùng nút “Cần sửa” để gửi phản hồi.")
    add_bullet(doc, "Bạn cũng có thể dùng lệnh /ask để gửi góp ý hoặc câu hỏi ẩn danh cho đội ngũ.")
    add_bullet(doc, "Với vấn đề về học liệu, email hoặc quyền truy cập Then trên web, liên hệ quản trị viên và cung cấp đúng email đã đăng ký.")

    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(closing, before=12, after=0, line=1.15)
    add_text(closing, "Chúc bạn học tập hiệu quả và khai thác hệ thống một cách chủ động.", size=10.5, color=MUTED, italic=True)

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
